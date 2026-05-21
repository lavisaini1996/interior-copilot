import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple


@dataclass
class QAItem:
    number: int
    question: str
    answer_lines: List[str]


QUESTION_RE = re.compile(r"^(?P<num>\d+)\.\s+###\s+(?P<q>.+?)\s*$")
BACK_TO_TOP_RE = re.compile(r"^\*\*\[⬆ Back to Top\]\(#table-of-contents\)\*\*\s*$")


def iter_lines(text: str) -> Iterable[str]:
    # Keep original line breaks; normalize to \n
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        yield line


def parse_qa(text: str) -> List[QAItem]:
    items: List[QAItem] = []
    current: Optional[QAItem] = None

    for line in iter_lines(text):
        m = QUESTION_RE.match(line)
        if m:
            if current:
                items.append(current)
            current = QAItem(
                number=int(m.group("num")),
                question=m.group("q").strip(),
                answer_lines=[],
            )
            continue

        if current is None:
            continue

        if BACK_TO_TOP_RE.match(line.strip()):
            # Drop boilerplate
            continue

        current.answer_lines.append(line)

    if current:
        items.append(current)

    # De-duplicate by question text (keep first occurrence)
    seen = set()
    deduped: List[QAItem] = []
    for it in items:
        key = re.sub(r"\s+", " ", it.question.strip().lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(it)

    return deduped


def split_blocks(lines: List[str]) -> List[Tuple[str, List[str]]]:
    """
    Split answer markdown-ish into blocks: ('code', [...]) or ('text', [...]).
    We only need to preserve fenced code blocks; everything else becomes paragraphs/bullets.
    """
    blocks: List[Tuple[str, List[str]]] = []
    in_code = False
    buf: List[str] = []

    def flush(kind: str):
        nonlocal buf
        if not buf:
            return
        # Trim leading/trailing empty lines for cleanliness
        while buf and buf[0].strip() == "":
            buf.pop(0)
        while buf and buf[-1].strip() == "":
            buf.pop()
        if buf:
            blocks.append((kind, buf))
        buf = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code:
                # closing fence
                flush("code")
                in_code = False
            else:
                flush("text")
                in_code = True
            continue

        buf.append(line)

    flush("code" if in_code else "text")
    return blocks


def is_bullet(line: str) -> bool:
    s = line.lstrip()
    return s.startswith("- ") or s.startswith("* ") or re.match(r"^\d+\.\s+", s) is not None


def clean_text_line(line: str) -> str:
    # Remove markdown emphasis markers but keep text.
    s = line.strip()
    s = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", s)  # links -> text
    s = s.replace("**", "").replace("__", "").replace("`", "")
    return s


def build_docx(items: List[QAItem], out_path: Path, title: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for idx, it in enumerate(items, start=1):
        if idx == 1 or idx % 25 == 0:
            print(f"Writing question {idx}/{len(items)}...", flush=True)

        doc.add_heading(f"{it.number}. {it.question}", level=2)

        blocks = split_blocks(it.answer_lines)
        for kind, blines in blocks:
            if kind == "code":
                # Preformatted block: one paragraph with monospaced font, preserve newlines.
                code = "\n".join(blines).rstrip()
                if not code.strip():
                    continue
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                continue

            # Text block: turn into paragraphs/bullets
            paragraph_buf: List[str] = []

            def flush_paragraph_buf():
                nonlocal paragraph_buf
                if not paragraph_buf:
                    return
                joined = " ".join(clean_text_line(x) for x in paragraph_buf if clean_text_line(x))
                joined = re.sub(r"\s+", " ", joined).strip()
                if joined:
                    doc.add_paragraph(joined)
                paragraph_buf = []

            for raw in blines:
                if raw.strip() == "":
                    flush_paragraph_buf()
                    continue

                if raw.strip().startswith(">"):
                    flush_paragraph_buf()
                    text = clean_text_line(raw.lstrip()[1:]).strip()
                    if text:
                        doc.add_paragraph(text)
                    continue

                if is_bullet(raw):
                    flush_paragraph_buf()
                    s = raw.lstrip()
                    s = re.sub(r"^[-*]\s+", "", s)
                    s = re.sub(r"^\d+\.\s+", "", s)
                    text = clean_text_line(s).strip()
                    if text:
                        doc.add_paragraph(text, style="List Bullet")
                    continue

                paragraph_buf.append(raw)

            flush_paragraph_buf()

        doc.add_paragraph()  # spacing between QAs

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    source = Path(r"c:\Users\Lavi\Downloads\# Nodejs Interview Questions and An.txt")
    out = Path(__file__).resolve().parents[1] / "Nodejs_Interview_QA.docx"

    text = source.read_text(encoding="utf-8", errors="replace")
    items = parse_qa(text)

    if not items:
        raise SystemExit("No questions found. Expected lines like: '12. ### <question>'")

    # The source file often contains multiple merged question sets, which can explode
    # into thousands of "unique" questions. By default, keep a practical core set.
    max_num = int(__import__("os").environ.get("MAX_Q", "300"))
    filtered = [it for it in items if it.number <= max_num]
    if not filtered:
        filtered = items

    print(
        f"Parsed {len(items)} unique questions; keeping {len(filtered)} (MAX_Q={max_num}). Building DOCX...",
        flush=True,
    )
    build_docx(
        items=filtered,
        out_path=out,
        title="Node.js Interview Questions & Answers",
    )
    print(f"Wrote: {out} ({len(filtered)} questions)")


if __name__ == "__main__":
    main()

