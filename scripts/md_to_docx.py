"""Convert docs/BLOG.md to docs/BLOG.docx — narrative blog, no tables."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "BLOG.md"
OUT_PATH = ROOT / "docs" / "BLOG.docx"


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            p.add_run(part)


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        else:
            add_rich_paragraph(doc, line.strip())

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    doc.save(out_path)


if __name__ == "__main__":
    md_to_docx(MD_PATH.read_text(encoding="utf-8"), OUT_PATH)
    print(f"Wrote {OUT_PATH}")
