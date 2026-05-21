from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_codeblock(doc: Document, text: str) -> None:
    for line in text.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def main() -> None:
    out_path = Path(__file__).resolve().parents[1] / "InteriorCopilot_ProblemSummary_Flow_ImplementationSteps.docx"

    doc = Document()

    title = doc.add_paragraph("Interior Copilot (Design Co-pilot) — Floor Plan → Catalog-Only Designs (2–3) with Priced BOM")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Date: {date.today().isoformat()}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # 1) Problem summary
    doc.add_heading("1. Problem summary (from transcript)", level=2)
    doc.add_paragraph(
        "Tilicho has launched an operations-focused app (system of record / SOR) and now wants a similar capability "
        "for the design function. The core pain is not creating mood boards (designers can do that quickly), but "
        "reducing the time and rework between initial requirements and the first usable design cut, while staying "
        "within Tilicho’s feasible ‘universe’ of materials/components (catalog)."
    )

    doc.add_paragraph("Key points captured:")
    for bullet in [
        "Mood boards are low value as an AI feature (designers can finish in ~30–40 minutes anyway).",
        "AI is useful if it helps after design starts: given dimensions + theme + color constraints, generate multiple viable options fast.",
        "Generic image generators (e.g., Gemini public image generation) are unreliable for constrained design; they inject creativity (random glass/handles) instead of obeying strict constraints.",
        "A major source of rework is selecting something that is not in Tilicho’s available catalog/universe; then the design must be changed, re-rendered, and the client must be re-convinced.",
        "Prompt enhancement and picking the right model are likely central (P0) to improving output quality.",
        "Cataloging + SOR for selections is still important for scale (P1/P2): prevent ad-hoc vendor knowledge living in individuals; ensure consistent dropdown selections; avoid appearing like a system integrator.",
        "Need example artifacts: questionnaires, SOR examples, estimates/first-cut artifacts, and end-to-end project docs to train/evaluate the co-pilot.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("")

    doc.add_heading("1.1 Updated MVP requirement (floor plan → catalog-only designs)", level=2)
    doc.add_paragraph(
        "Replace ‘moodboard’ input with a floor plan input (image or text). The co-pilot should parse/understand the "
        "floor plan, ask the user targeted questions when required information is missing or ambiguous (style references, "
        "materials/finishes, budget bands, constraints), then generate 2–3 design options. Every option must use ONLY items "
        "from Tilicho’s catalog (no invented components). Output should include generated images for the options and a "
        "priced materials/BOM list per option."
    )

    doc.add_paragraph("Non-negotiable constraints for this MVP:")
    for bullet in [
        "Input: accept floor plan as image (e.g., JPG/PNG/PDF) or structured text dimensions.",
        "Clarification: if key info is missing, ask questions; if user doesn’t specify, use safe defaults (and clearly list assumptions).",
        "Generation: create 2–3 options, each strictly catalog-only.",
        "Output: images + priced material list (catalog item name/ID, qty, unit price, subtotal, total).",
        "Governance: never place an item that is not in catalog; instead propose substitutions from catalog.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("")

    # 2) Proposed solution framing
    doc.add_heading("2. Proposed solution framing", level=2)
    doc.add_paragraph(
        "Build a Design Co-pilot that extends the System of Record into the design workflow and focuses on constraint-following, "
        "catalog-aware generation. The co-pilot should help designers turn client conversations and structured inputs into: "
        "(a) clean constraints, (b) better prompts/briefs, (c) 2–3 floor-plan-aware first-cut options that remain inside the catalog universe, "
        "and (d) fewer iteration cycles caused by unavailable selections — while producing a priced BOM for every option."
    )

    doc.add_paragraph("What success looks like (high-level outcomes):")
    for bullet in [
        "Faster first-cut: reduce time from requirements → first design cut.",
        "Higher constraint adherence: outputs follow dimensions, layout rules, and style constraints.",
        "Catalog compliance: reduce designs that contain unavailable materials/components.",
        "Lower rework: fewer redesign cycles due to feasibility mismatches.",
        "Scalable process: new designers ramp faster with consistent internal universe + SOR data.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph("")

    # 3) Flowchart (text-based, stable in docx)
    doc.add_heading("3. End-to-end flowchart", level=2)
    doc.add_paragraph(
        "Flowchart is provided as a fixed-width text diagram so it renders reliably in .docx across machines."
    )

    flow = r"""
┌───────────────────────────────────────┐
│ Input: Floor plan                     │
│ - Image (JPG/PNG/PDF) OR              │
│ - Text (room dims / walls / openings) │
└───────────────────┬───────────────────┘
                    │
                    ▼
┌───────────────────────────────────────┐
│ Floor plan understanding              │
│ - detect rooms/zones (if possible)    │
│ - extract key dims & constraints      │
│ - note unknowns (doors/windows)       │
└───────────────────┬───────────────────┘
                    │
                    ▼
┌───────────────────────────────────────┐
│ SOR capture (design brief)            │
│ - functional needs                    │
│ - budgets (bands)                     │
│ - style references                    │
│ - do/don't rules                      │
└───────────────────┬───────────────────┘
                    │
                    ▼
┌───────────────────────────────────────┐
│ Missing / ambiguous info?             │
└───────────────┬───────────────┬───────┘
                │Yes            │No
                ▼               ▼
┌───────────────────────────┐  ┌───────────────────────────┐
│ Ask user targeted         │  │ Normalize constraints      │
│ questions                 │  │ - units & tolerances       │
│ - style/material          │  │ - must-have vs optional    │
│ - budget band             │  │ - safe defaults (recorded) │
│ - constraints/priority    │  └───────────────┬───────────┘
└───────────────┬───────────┘                  │
                │                               ▼
                └───────────────►┌───────────────────────────┐
                                 │ Catalog gating            │
                                 │ - select ONLY catalog IDs │
                                 │ - substitutions if needed │
                                 └───────────────┬───────────┘
                                                 │
                                                 ▼
                                 ┌───────────────────────────┐
                                 │ Generate 2–3 options       │
                                 │ - layout + placements      │
                                 │ - strict catalog-only      │
                                 └───────────────┬───────────┘
                                                 │
                                                 ▼
                                 ┌───────────────────────────┐
                                 │ Output package per option  │
                                 │ - images (generated)       │
                                 │ - priced BOM/material list │
                                 │ - assumptions              │
                                 └───────────────┬───────────┘
                                                 │
                                                 ▼
                                 ┌───────────────────────────┐
                                 │ Iterate (feedback → SOR)   │
                                 │ - regenerate deltas only   │
                                 └───────────────────────────┘
""".strip("\n")
    _add_codeblock(doc, flow)

    doc.add_paragraph("")

    # 4) Implementation steps (phased)
    doc.add_heading("4. Steps of implementation (phased)", level=2)

    doc.add_heading("Phase 0 — Inputs & baseline (1–2 weeks)", level=3)
    for bullet in [
        "Collect 1–2 end-to-end project samples: questionnaire, site measurements, mood board (optional), first-cut estimate, final design outputs, and iterations.",
        "Define the SOR schema for design: rooms, modules (e.g., wardrobe), dimensions, constraints, materials, catalog SKUs, and decision history.",
        "Define evaluation rubric: constraint adherence, catalog compliance, first-cut speed, number of iteration cycles, designer satisfaction.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("Phase 1 — SOR extension + catalog foundation (P0/P1)", level=3)
    for bullet in [
        "Implement design SOR capture UI/API: structured forms for dimensions, style, do/don’t rules, and client preferences.",
        "Create catalog service: normalized materials/components/colors/finishes with IDs, availability, and substitution rules.",
        "Add feasibility validator: detect out-of-universe selections early and suggest catalog-safe alternatives.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("Phase 2 — Prompt enhancement engine (P0)", level=3)
    for bullet in [
        "Build a constraint-to-prompt compiler: transforms SOR + catalog constraints into a strict, model-friendly brief.",
        "Add prompt guardrails: explicit do/don’t lists, geometry constraints, and ‘no extra elements’ enforcement.",
        "Model selection & routing: test candidate models; route tasks (layout vs styling vs copy) to best-fit model.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("Phase 3 — First-cut generation & packaging (P0)", level=3)
    for bullet in [
        "Generate multiple options (N variants) with controlled randomness only within allowed parameters.",
        "Produce a first-cut package: images/visuals + annotated dimensions + bill of materials tied to catalog IDs.",
        "Provide ‘delta regeneration’: when feedback changes one constraint, regenerate only impacted parts.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("Phase 4 — Designer workflow integration (P0)", level=3)
    for bullet in [
        "Designer review UI: side-by-side options, constraint diffs, quick edits to SOR, and regenerate button.",
        "Approval + traceability: every output linked to SOR version and catalog snapshot for auditability.",
        "Export/handoff: outputs and selections flow into estimating/operations modules.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_heading("Phase 5 — Scale & governance (P1/P2)", level=3)
    for bullet in [
        "Knowledge centralization: reduce dependency on individual vendor/shop knowledge by enforcing catalog-first selections.",
        "Quality monitoring: track invalid generations, catalog violations, and feedback reasons to improve prompts and rules.",
        "Role-based access and templates: room/module templates, approved styles, and standard constraint presets.",
    ]:
        doc.add_paragraph(bullet, style="List Number")

    doc.add_paragraph("")

    doc.add_heading("5. Open questions / dependencies", level=2)
    for bullet in [
        "What design outputs are required for ‘first cut’ (2D layout, 3D render, elevations, BOM, estimate)?",
        "What is the current catalog source of truth and how often does it change?",
        "What tools do designers use today (CAD/SketchUp/others) for integration and file formats?",
        "What are the most common rework triggers (color mismatch, dimensions, budget, availability, client preference changes)?",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()

